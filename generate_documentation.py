#!/usr/bin/env python3
"""
Generate Activity Documentation for QRentry Event Registration System
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_documentation():
    """Create comprehensive Activity Documentation"""
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_heading('QRentry Event Registration System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format = title.runs[0].font
    title_format.color.rgb = RGBColor(99, 102, 241)  # Primary color
    
    # Subtitle
    subtitle = doc.add_paragraph('Activity Documentation Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.italic = True
    subtitle_format.size = Pt(12)
    
    # Date
    date_para = doc.add_paragraph('January 11, 2026')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(10)
    date_para.runs[0].font.color.rgb = RGBColor(107, 114, 128)
    
    doc.add_paragraph()  # Spacing
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Project Overview',
        '2. System Description',
        '3. Features',
        '4. Architecture Overview',
        '5. Application Pages & Screenshots',
        '6. API Integration',
        '7. Installation & Setup Instructions',
        '8. Running the Project',
        '9. Technical Stack',
        '10. Support & Documentation'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 1. Project Overview
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        'QRentry is a comprehensive Event Registration System that streamlines event '
        'management through QR code technology. It enables attendees to register for events, '
        'organizers to manage events and registrations, and administrators to oversee the entire system.'
    )
    
    # 2. System Description
    doc.add_heading('2. System Description', level=1)
    doc.add_paragraph(
        'QRentry simplifies the event registration process by leveraging QR codes for '
        'quick check-ins, ticket management, and attendee tracking. The system supports multiple '
        'user roles (Attendee, Organizer, Admin) with different permissions and capabilities.'
    )
    
    description_table = doc.add_table(rows=5, cols=2)
    description_table.style = 'Light Grid Accent 1'
    
    header_cells = description_table.rows[0].cells
    header_cells[0].text = 'Aspect'
    header_cells[1].text = 'Details'
    
    rows_data = [
        ('System Name', 'QRentry Event Registration System'),
        ('Primary Purpose', 'Event registration, ticketing, and check-in management'),
        ('Target Users', 'Event Attendees, Event Organizers, System Administrators'),
        ('Key Technology', 'QR Code scanning for registration and check-in')
    ]
    
    for i, (aspect, detail) in enumerate(rows_data, 1):
        cells = description_table.rows[i].cells
        cells[0].text = aspect
        cells[1].text = detail
    
    doc.add_page_break()
    
    # 3. Features
    doc.add_heading('3. Features', level=1)
    
    doc.add_heading('Authentication', level=2)
    features_auth = [
        'User Registration (Attendee, Organizer, Admin)',
        'Secure Login with JWT authentication',
        'Password reset functionality',
        'Email verification',
        'Session management'
    ]
    for feature in features_auth:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('Event Management', level=2)
    features_event = [
        'Create and manage events',
        'Set event details, date, time, and location',
        'Manage event capacity',
        'View event registrations',
        'Generate QR codes for events'
    ]
    for feature in features_event:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('Registration & Ticketing', level=2)
    features_reg = [
        'Browse and register for events',
        'Download event tickets',
        'QR code integration for ticket verification',
        'Multiple registration categories',
        'Event status tracking'
    ]
    for feature in features_reg:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('Check-in Management', level=2)
    features_checkin = [
        'QR code scanning for attendee check-in',
        'Real-time attendance tracking',
        'Check-in status updates',
        'Attendee list management',
        'Check-in history and reporting'
    ]
    for feature in features_checkin:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('Admin Dashboard', level=2)
    features_admin = [
        'System overview and analytics',
        'User management',
        'Event oversight',
        'Registration monitoring',
        'System configuration'
    ]
    for feature in features_admin:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # 4. Architecture Overview
    doc.add_heading('4. Architecture Overview', level=1)
    
    doc.add_heading('Frontend Stack', level=2)
    doc.add_paragraph('React with TypeScript for type-safe component development', style='List Bullet')
    doc.add_paragraph('Tailwind CSS for responsive, utility-first styling', style='List Bullet')
    doc.add_paragraph('React Router for client-side navigation', style='List Bullet')
    doc.add_paragraph('Zustand for state management', style='List Bullet')
    doc.add_paragraph('Axios for HTTP API communication', style='List Bullet')
    doc.add_paragraph('QR code scanning library for ticket verification', style='List Bullet')
    
    doc.add_heading('Backend Stack', level=2)
    doc.add_paragraph('NestJS framework for robust API development', style='List Bullet')
    doc.add_paragraph('PostgreSQL database for data persistence', style='List Bullet')
    doc.add_paragraph('JWT (JSON Web Tokens) for authentication', style='List Bullet')
    doc.add_paragraph('TypeORM for database ORM', style='List Bullet')
    doc.add_paragraph('Role-based access control (RBAC)', style='List Bullet')
    doc.add_paragraph('RESTful API endpoints', style='List Bullet')
    
    doc.add_heading('Authentication Flow', level=2)
    doc.add_paragraph('1. User registers with email, password, and role (Attendee/Organizer/Admin)')
    doc.add_paragraph('2. User logs in with credentials')
    doc.add_paragraph('3. Backend validates and issues JWT token')
    doc.add_paragraph('4. Frontend stores token in secure storage')
    doc.add_paragraph('5. Token included in all subsequent API requests')
    doc.add_paragraph('6. Backend validates token on protected routes')
    
    doc.add_page_break()
    
    # 5. Application Pages & Screenshots
    doc.add_heading('5. Application Pages & Screenshots', level=1)
    doc.add_paragraph(
        'Note: The following pages represent the key user interfaces of QRentry. '
        'Screenshots should be captured from the running application.'
    )
    
    pages = [
        {
            'name': 'Login Page',
            'description': 'User authentication interface. Attendees, organizers, and admins use this page to log in with their email/username and password. Includes forgot password link and signup option.',
            'users': 'All user types'
        },
        {
            'name': 'Registration Page',
            'description': 'New user account creation. Users select their role (Attendee, Organizer, Admin), enter email, password, name, and optional company information.',
            'users': 'New users'
        },
        {
            'name': 'Discover Events Page',
            'description': 'Browse all available events in the system. Displays event cards with name, date, time, location, and registration button. Users can search and filter events.',
            'users': 'Attendees, Organizers'
        },
        {
            'name': 'Create Event Page',
            'description': 'Form for organizers to create new events. Includes fields for event name, description, date, time, location, capacity, and other event details.',
            'users': 'Organizers'
        },
        {
            'name': 'Event Details Page',
            'description': 'Detailed view of a single event. Shows full event information, registration status, attendee count, and action buttons (Register, View Tickets, Check-in).',
            'users': 'Attendees, Organizers'
        },
        {
            'name': 'My Events Page',
            'description': 'Dashboard showing registered events for attendees. Displays list of events with registration status, allows filtering by status (upcoming, past, cancelled).',
            'users': 'Attendees'
        },
        {
            'name': 'My Tickets Page',
            'description': 'View and manage event tickets. Shows QR codes for each ticket, ticket details, and options to download or print tickets. Tickets display registration information.',
            'users': 'Attendees'
        },
        {
            'name': 'Check-in Page',
            'description': 'QR code scanner interface for checking in attendees. Organizers and admins can scan attendee QR codes to verify attendance. Shows real-time attendance statistics.',
            'users': 'Organizers, Admins'
        },
        {
            'name': 'Admin Dashboard',
            'description': 'System overview for administrators. Shows key metrics including total events, registrations, users, and system status. Provides access to manage users, events, and registrations.',
            'users': 'Admins'
        },
        {
            'name': 'Admin Users Management',
            'description': 'Manage system users. Admins can view user list, edit user details, change user roles, and delete accounts. Shows user registration date and activity status.',
            'users': 'Admins'
        }
    ]
    
    for i, page in enumerate(pages, 1):
        doc.add_heading(f'{i}. {page["name"]}', level=2)
        
        desc_para = doc.add_paragraph()
        desc_para.add_run('Description: ').bold = True
        desc_para.add_run(page['description'])
        
        users_para = doc.add_paragraph()
        users_para.add_run('Accessible to: ').bold = True
        users_para.add_run(page['users'])
        
        # Placeholder for screenshot
        screenshot_note = doc.add_paragraph()
        screenshot_note.add_run('[Screenshot of ' + page['name'] + ']').italic = True
        screenshot_note.runs[0].font.color.rgb = RGBColor(107, 114, 128)
        
        doc.add_paragraph()  # Spacing
    
    doc.add_page_break()
    
    # 6. API Integration
    doc.add_heading('6. API Integration', level=1)
    doc.add_paragraph('QRentry frontend communicates with the backend API for all operations.')
    
    api_table = doc.add_table(rows=1, cols=4)
    api_table.style = 'Light Grid Accent 1'
    
    hdr_cells = api_table.rows[0].cells
    hdr_cells[0].text = 'Endpoint'
    hdr_cells[1].text = 'Method'
    hdr_cells[2].text = 'Purpose'
    hdr_cells[3].text = 'Auth Required'
    
    endpoints = [
        ('POST /api/auth/register', 'POST', 'User registration', 'No'),
        ('POST /api/auth/login', 'POST', 'User authentication', 'No'),
        ('GET /api/events', 'GET', 'List all events', 'No'),
        ('POST /api/events', 'POST', 'Create new event', 'Yes'),
        ('GET /api/events/:id', 'GET', 'Get event details', 'No'),
        ('POST /api/registrations', 'POST', 'Register for event', 'Yes'),
        ('GET /api/my-registrations', 'GET', 'User registrations', 'Yes'),
        ('POST /api/check-in', 'POST', 'Check in attendee', 'Yes'),
        ('GET /api/admin/dashboard', 'GET', 'Admin stats', 'Yes (Admin)'),
        ('GET /api/admin/users', 'GET', 'Manage users', 'Yes (Admin)')
    ]
    
    for endpoint, method, purpose, auth in endpoints:
        row_cells = api_table.add_row().cells
        row_cells[0].text = endpoint
        row_cells[1].text = method
        row_cells[2].text = purpose
        row_cells[3].text = auth
    
    doc.add_page_break()
    
    # 7. Installation & Setup
    doc.add_heading('7. Installation & Setup Instructions', level=1)
    
    doc.add_heading('Prerequisites', level=2)
    doc.add_paragraph('Node.js (v16 or higher)', style='List Bullet')
    doc.add_paragraph('npm or yarn package manager', style='List Bullet')
    doc.add_paragraph('PostgreSQL database (v12 or higher)', style='List Bullet')
    doc.add_paragraph('Git for version control', style='List Bullet')
    
    doc.add_heading('Step 1: Clone Repository', level=2)
    code1 = doc.add_paragraph('git clone <repository-url>')
    code1.style = 'No Spacing'
    code1.runs[0].font.name = 'Courier New'
    code1.runs[0].font.color.rgb = RGBColor(31, 41, 55)
    
    code2 = doc.add_paragraph('cd event-registration-system')
    code2.style = 'No Spacing'
    code2.runs[0].font.name = 'Courier New'
    code2.runs[0].font.color.rgb = RGBColor(31, 41, 55)
    
    doc.add_heading('Step 2: Backend Setup', level=2)
    doc.add_paragraph('Navigate to backend directory')
    code3 = doc.add_paragraph('cd backend')
    code3.style = 'No Spacing'
    code3.runs[0].font.name = 'Courier New'
    
    doc.add_paragraph('Install dependencies')
    code4 = doc.add_paragraph('npm install')
    code4.style = 'No Spacing'
    code4.runs[0].font.name = 'Courier New'
    
    doc.add_paragraph('Create .env file with database configuration')
    doc.add_paragraph('Run database migrations (if applicable)')
    code5 = doc.add_paragraph('npm run typeorm migration:run')
    code5.style = 'No Spacing'
    code5.runs[0].font.name = 'Courier New'
    
    doc.add_heading('Step 3: Frontend Setup', level=2)
    doc.add_paragraph('Navigate to frontend directory')
    code6 = doc.add_paragraph('cd frontend')
    code6.style = 'No Spacing'
    code6.runs[0].font.name = 'Courier New'
    
    doc.add_paragraph('Install dependencies')
    code7 = doc.add_paragraph('npm install')
    code7.style = 'No Spacing'
    code7.runs[0].font.name = 'Courier New'
    
    doc.add_paragraph('Configure API base URL in .env or config files')
    
    doc.add_page_break()
    
    # 8. Running the Project
    doc.add_heading('8. Running the Project', level=1)
    
    doc.add_heading('Backend Startup', level=2)
    doc.add_paragraph('From backend directory:')
    code_back = doc.add_paragraph('npm run start:dev')
    code_back.style = 'No Spacing'
    code_back.runs[0].font.name = 'Courier New'
    doc.add_paragraph('Backend runs on: http://localhost:3001')
    
    doc.add_heading('Frontend Startup', level=2)
    doc.add_paragraph('From frontend directory (in separate terminal):')
    code_front = doc.add_paragraph('npm start')
    code_front.style = 'No Spacing'
    code_front.runs[0].font.name = 'Courier New'
    doc.add_paragraph('Frontend runs on: http://localhost:3000')
    
    doc.add_heading('Accessing the Application', level=2)
    doc.add_paragraph('Open browser and navigate to: http://localhost:3000')
    doc.add_paragraph('Create a new account or use provided test credentials')
    doc.add_paragraph('Select role: Attendee, Organizer, or Admin')
    
    doc.add_heading('Test Credentials (if available)', level=2)
    test_table = doc.add_table(rows=4, cols=3)
    test_table.style = 'Light Grid Accent 1'
    
    test_hdr = test_table.rows[0].cells
    test_hdr[0].text = 'Role'
    test_hdr[1].text = 'Email'
    test_hdr[2].text = 'Password'
    
    test_data = [
        ('Admin', 'admin@test.com', 'Test@123'),
        ('Organizer', 'organizer@test.com', 'Test@123'),
        ('Attendee', 'attendee@test.com', 'Test@123')
    ]
    
    for role, email, pwd in test_data:
        row = test_table.add_row()
        row.cells[0].text = role
        row.cells[1].text = email
        row.cells[2].text = pwd
    
    doc.add_page_break()
    
    # 9. Technical Stack
    doc.add_heading('9. Technical Stack', level=1)
    
    doc.add_heading('Frontend Technologies', level=2)
    stack_front = [
        'React 18.x - UI framework',
        'TypeScript - Type-safe JavaScript',
        'Tailwind CSS - Utility-first CSS framework',
        'React Router - Navigation',
        'Zustand - State management',
        'Axios - HTTP client',
        'QR Code Scanner - QR code reading'
    ]
    for tech in stack_front:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_heading('Backend Technologies', level=2)
    stack_back = [
        'NestJS - Progressive Node.js framework',
        'TypeScript - Type-safe JavaScript',
        'PostgreSQL - Relational database',
        'TypeORM - ORM for database operations',
        'JWT - Authentication tokens',
        'bcrypt - Password hashing',
        'Passport - Authentication middleware'
    ]
    for tech in stack_back:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_heading('Development Tools', level=2)
    stack_dev = [
        'VS Code - Code editor',
        'Git - Version control',
        'npm - Package manager',
        'Postman - API testing (optional)',
        'Docker - Containerization (optional)'
    ]
    for tech in stack_dev:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_page_break()
    
    # 10. Support & Documentation
    doc.add_heading('10. Support & Documentation', level=1)
    
    doc.add_heading('Troubleshooting', level=2)
    doc.add_paragraph('Port Already in Use: Change port in configuration or kill process using the port', style='List Bullet')
    doc.add_paragraph('Database Connection Issues: Verify PostgreSQL is running and .env credentials are correct', style='List Bullet')
    doc.add_paragraph('Module Not Found: Run npm install again in respective directory', style='List Bullet')
    doc.add_paragraph('CORS Errors: Ensure backend CORS configuration includes frontend URL', style='List Bullet')
    doc.add_paragraph('Authentication Failures: Check JWT token expiration and localStorage/sessionStorage settings', style='List Bullet')
    
    doc.add_heading('Common Commands', level=2)
    commands = [
        ('npm install', 'Install project dependencies'),
        ('npm start', 'Start frontend development server'),
        ('npm run start:dev', 'Start backend in development mode'),
        ('npm run build', 'Build frontend for production'),
        ('npm run test', 'Run tests'),
        ('npm run lint', 'Run code linter'),
    ]
    
    cmd_table = doc.add_table(rows=1, cols=2)
    cmd_table.style = 'Light Grid Accent 1'
    
    cmd_hdr = cmd_table.rows[0].cells
    cmd_hdr[0].text = 'Command'
    cmd_hdr[1].text = 'Description'
    
    for cmd, desc in commands:
        cmd_row = cmd_table.add_row()
        cmd_row.cells[0].text = cmd
        cmd_row.cells[1].text = desc
    
    doc.add_heading('Resources', level=2)
    doc.add_paragraph('React Documentation: https://react.dev', style='List Bullet')
    doc.add_paragraph('NestJS Documentation: https://docs.nestjs.com', style='List Bullet')
    doc.add_paragraph('Tailwind CSS: https://tailwindcss.com', style='List Bullet')
    doc.add_paragraph('TypeScript: https://www.typescriptlang.org', style='List Bullet')
    doc.add_paragraph('PostgreSQL: https://www.postgresql.org', style='List Bullet')
    
    doc.add_page_break()
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'QRentry is a comprehensive event registration and management system that combines '
        'modern web technologies with practical event management features. The system is designed '
        'to be user-friendly, scalable, and secure. By following the setup and running instructions, '
        'you should have a fully functional event management platform ready to deploy.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph('For additional support, refer to the README files in the respective directories.')
    
    # Save document
    output_path = 'QRentry_Activity_Documentation.docx'
    doc.save(output_path)
    
    print(f'✅ Documentation created successfully: {output_path}')
    return output_path

if __name__ == '__main__':
    create_documentation()
